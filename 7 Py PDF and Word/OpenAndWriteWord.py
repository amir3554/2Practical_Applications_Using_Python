import docx
from pathlib import Path
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx.styles

mydoc = docx.Document()

mydoc.add_paragraph("This is the first paragraph.")
mydoc.add_paragraph("This is the second paragraph.")

para = mydoc.add_paragraph("هذه هي الفقرة الثالثة ولو.")
para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

forth_para = mydoc.add_paragraph("This is the forth paragraph...")
forth_para.add_run("and this is a run .")

mydoc.add_heading("THIS IS HEADING , LEVEL 1", 0)
mydoc.add_heading("THIS IS HEADING , LEVEL 2", 1)
mydoc.add_heading("THIS IS HEADING , LEVEL 3", 2)
mydoc.add_heading("THIS IS HEADING , LEVEL 4", 3)
mydoc.add_heading("THIS IS HEADING , LEVEL 5", 4)

print(mydoc.paragraphs[0].style)
print(mydoc.paragraphs[4].style)
print(mydoc.paragraphs[5].style)
mydoc.paragraphs[0].style = mydoc.styles["Heading 1"]
mydoc.paragraphs[3].style = mydoc.styles["Heading 1"]
mydoc.paragraphs[3].style.delete()

mydoc.add_paragraph("This is Styled paragrahp", "Heading 4")

mydoc.add_picture(str(Path.home() / Path("Desktop", "tests",)))

mydoc.save(str(Path.home() / Path("Desktop", "tests", "From Python.docx")))



