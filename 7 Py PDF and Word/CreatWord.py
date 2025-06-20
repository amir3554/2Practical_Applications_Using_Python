import docx
from pathlib import Path
import DefReadWord

doc = docx.Document(Path.home() / Path('Desktop', 'tests', 'academy_1.docx'))

print(len(doc.paragraphs))

print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(doc.paragraphs[2].text)

print(doc.paragraphs[2].runs[0].text)
print(doc.paragraphs[2].runs[1].text)
print(doc.paragraphs[2].runs[2].text)
print(doc.paragraphs[2].runs[3].text)
print(doc.paragraphs[2].runs[4].text)

print("************************************************")

print(DefReadWord.get_text(Path.home() / Path('Desktop', 'tests', 'academy_1.docx')))


