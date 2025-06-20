# Making a program that reminds the members who hasn't paid for their membership at the Gym .

def pay_check(file_name :str) -> list:
    import openpyxl
    from pathlib import Path
    file_excil = openpyxl.load_workbook(Path.home() / Path("Desktop", "tests", f"{str(file_name)}"))
    The_Sheet = file_excil['Sheet1']


    The_cols = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I']
    The_rows = list(range(The_Sheet.max_row +1))
    All_data_list = []

    for row in The_rows[2:]:
        The_single_row_list = []
        ids = str(The_Sheet[f'A{row}'].value)
        names = The_Sheet[f'B{row}'].value
        emails = The_Sheet[f'C{row}'].value
        dates_not_paid = []
        for col in The_cols[3:]:
            if not The_Sheet[f'{col}{row}'].value:
                dates_not_paid.append(str(The_Sheet[f'{col}1'].value))
        The_single_row_list.append(ids)
        The_single_row_list.append(names)
        The_single_row_list.append(emails)
        The_single_row_list.append(dates_not_paid)
        All_data_list.append(The_single_row_list)
        
    return All_data_list


def the_reminder_docx(The_data_list : list):
    from docx import Document
    from pathlib import Path
    for Num, data in enumerate(The_data_list):
        The_doc = Document()
        if not data[3]:
            continue
        #[ID = '7',   Name = "amir",   Email = "amir@example.com",   Times not paid = 0]
        The_doc.add_heading(f"Hello {data[1]}", 0)
        The_doc.add_paragraph(f"You haven't paid your Gym membership for {len(data[3])} times, Here are the details:({' '.join(map(str, data[3]))}) times, Your ID is {data[0]}, Please reply on this email {data[2]}.")
        The_doc.save(Path.home() / Path("Desktop", "tests", f"alquiler{Num}_{data[1]}_{data[0]}.docx"))


def the_reminder_SMTP(The_data_list : list):
    import smtplib
    SMTP_Obj = smtplib.SMTP("smtp.example.com", 587)
    SMTP_Obj.starttls()
    email_sender = "amirdwikat@example.com"
    password = str(input("Please inter the password. "))
    SMTP_Obj.login(email_sender, password=password)

    for data in The_data_list:
        The_Body =f"SUBJECT: Dear {data[1]}\nRecords show that you haven't paid for {len(data[3])} times, Here are the details:({' '.join(map(str, data[3]))}), Please make this paiment as soon as posible, Thank you."
        if not data[3]:
            continue
        email_rec = data[2] 
        SMTP_Obj.sendmail(email_sender, email_rec, The_Body)
    SMTP_Obj.quit()


try:
    the_reminder_docx(pay_check("members.xlsx"))
    pass


except Exception as e:
    print(f"This is an ERROR: -->\n{e}")
